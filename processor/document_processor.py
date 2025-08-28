import os
import PyPDF2
import docx

class DocumentProcessor:
    """Handles extraction of text content from PDF and DOCX files"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting PDF content: {e}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting DOCX content: {e}")
    
    @staticmethod
    def extract_content(file_path: str) -> str:
        """Extract content based on file extension"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return DocumentProcessor.extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return DocumentProcessor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported: .pdf, .docx, .doc")
